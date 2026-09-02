from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUT = DOCS / "Mahosample-final-project-report.docx"

SOURCE_FILES = [
    ("รายงานฉบับสมบูรณ์", DOCS / "final-project-report.md"),
    ("ภาคผนวก A: SRS", DOCS / "srs.md"),
    ("ภาคผนวก B: RTM", DOCS / "rtm.md"),
    ("ภาคผนวก C: Architecture, ER/Data Model", DOCS / "architecture.md"),
    ("ภาคผนวก D: API Contract", DOCS / "api-contract.md"),
    ("ภาคผนวก E: AI Usage Log", DOCS / "ai-usage-log.md"),
    ("ภาคผนวก F: Production Cleanup Evidence", DOCS / "production-cleanup-2026-09-02.md"),
    ("ภาคผนวก G: Test Cases", DOCS / "test-cases.md"),
    ("ภาคผนวก H: Test Report", DOCS / "test-report-2026-09-02.md"),
    ("ภาคผนวก I: Staff User Guide", DOCS / "staff-user-guide.md"),
]


def set_run_font(run, size=None, bold=None, color=None, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for edge in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, edge, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    set_run_font(footer.add_run("Mahosample Final Project Report | Page "), size=9, color="666666")
    add_page_number(footer)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Mahosample")
    set_run_font(r, size=28, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Final Software Project Report")
    set_run_font(r, size=18, color="2E74B5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ระบบลงทะเบียนรับตัวอย่าง มะโฮ และจัดการข้อมูลจัดส่ง")
    set_run_font(r, size=14, color="333333")

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    set_table_width(table, [1.8, 4.6])
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, "F7F8FA")
    rows = [
        ("Project", "Mahosample"),
        ("Live URL", "http://maho.kitaith.com:18080/"),
        ("Repository", "https://github.com/kit-sinlapasa/Mahosample"),
        ("Date", "2026-09-02"),
        ("Status", "ผ่านการทดสอบระบบหลักและจัดทำเอกสารส่งงานครบชุด"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(cell is row.cells[0]))

    doc.add_page_break()


def add_toc_placeholder(doc):
    h = doc.add_paragraph("สารบัญ", style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph()
    r = p.add_run("สามารถคลิกขวาที่สารบัญใน Microsoft Word แล้วเลือก Update Field เพื่ออัปเดตเลขหน้าอัตโนมัติ")
    set_run_font(r, size=10, color="666666")

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "คลิกขวาเพื่อ Update Field"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    doc.add_page_break()


def split_markdown_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        line = lines[i].strip()
        cells = [c.strip() for c in line.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_text(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, font="Consolas", size=10, color="333333")
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    total_width = 6.5
    widths = [total_width / col_count] * col_count
    set_table_width(table, widths)
    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            value = row[col_idx] if col_idx < len(row) else ""
            cell.text = re.sub(r"`([^`]+)`", r"\1", value)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run_font(run, size=9.2, bold=(row_idx == 0))


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(code.rstrip())
    set_run_font(r, font="Consolas", size=8.5, color="333333")


def add_markdown_file(doc, title, path, first=False):
    if not first:
        doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph(title, style="Heading 1")
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            rows, i = split_markdown_table(lines, i)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            doc.add_paragraph(heading.group(2).strip(), style=f"Heading {level}")
            i += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_markdown_text(p, bullet.group(1))
            i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_markdown_text(p, numbered.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_markdown_text(p, line)
        i += 1

    if in_code and code_lines:
        add_code_block(doc, "\n".join(code_lines))


def main():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_toc_placeholder(doc)
    for idx, (title, path) in enumerate(SOURCE_FILES):
        add_markdown_file(doc, title, path, first=(idx == 0))
    doc.core_properties.title = "Mahosample Final Project Report"
    doc.core_properties.subject = "Final Software Project Report"
    doc.core_properties.author = "Mahosample Project Team"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
