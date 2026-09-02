from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from build_final_docx import set_cell_shading, set_run_font, set_table_width, setup_styles
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
EVIDENCE = ROOT / "outputs" / "e2e-evidence"
SCREENSHOTS = EVIDENCE / "screenshots"
DOCX_IMAGE_DIR = EVIDENCE / "docx-images"
OUTPUT = DOCS / "Mahosample-test-evidence-with-screenshots.docx"

SCREENSHOT_LABELS = {
    "public-registration-form": ("TC-PUB-001", "หน้าแบบฟอร์มลงทะเบียนรับตัวอย่าง มะโฮ"),
    "public-contact-validation": ("TC-PUB-002", "Validation เมื่อไม่กรอกอีเมล, LINE ID หรือ Facebook Messenger"),
    "public-success-card": ("TC-PUB-001", "Success card หลังส่งฟอร์ม พร้อมเลขรายการภายใน"),
    "tracking-found": ("TC-TRK-001", "หน้าเช็คสถานะพบเลข tracking และ Tracking URL"),
    "tracking-not-found": ("TC-TRK-002", "หน้าเช็คสถานะเมื่อไม่พบเลข tracking"),
    "staff-dashboard": ("TC-STF-003", "Dashboard พนักงานหลัง login"),
    "staff-full-table": ("TC-STF-006", "หน้าตารางข้อมูลเต็ม แสดง column ข้อมูลทั้งหมด"),
    "staff-filter-province": ("TC-STF-008", "Filter header ตารางด้วยจังหวัด"),
    "staff-date-range-filter": ("TC-STF-009/010", "Filter วันที่ส่งและวันที่ลงทะเบียนแบบช่วงเวลา"),
    "staff-sorting": ("TC-STF-007", "Sorting จาก header ตาราง"),
    "staff-detail-drawer": ("TC-STF-004/005/006", "Drawer รายละเอียด พร้อมบันทึกคำขอและอัปเดตขนส่ง"),
    "staff-tracking-url-link": ("TC-STF-012", "Tracking URL แสดงเป็น link คลิกได้"),
    "staff-export": ("TC-EXP-001/002", "Export CSV จากหน้าพนักงาน"),
}


def latest_json() -> Path:
    files = sorted(EVIDENCE.glob("live-e2e-results-*.json"), key=lambda p: p.stat().st_mtime)
    if not files:
        raise FileNotFoundError("No live e2e result JSON found")
    return files[-1]


def latest_screenshot_stamp() -> str:
    stamps = []
    for file in SCREENSHOTS.glob("*.png"):
        match = re.match(r"(\d{14})-", file.name)
        if match:
            stamps.append(match.group(1))
    if not stamps:
        raise FileNotFoundError("No screenshot files found")
    return sorted(stamps)[-1]


def add_cover(doc, result):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Mahosample")
    set_run_font(r, size=28, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Test Evidence With Screenshots")
    set_run_font(r, size=18, color="2E74B5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("รายงานผลการทดสอบระบบพร้อมภาพหน้าจอประกอบ")
    set_run_font(r, size=14, color="333333")

    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=2)
    set_table_width(table, [2.0, 4.3])
    rows = [
        ("System", "Mahosample"),
        ("Environment", result["base_url"]),
        ("Run ID", result["run_id"]),
        ("Executed At", result["executed_at"]),
        ("Result", f'{result["passed"]}/{result["total"]} passed, {result["failed"]} failed'),
        ("Scope", "Public registration, tracking, staff dashboard, full table, export/import, permission"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "F2F4F7")
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(cell is row.cells[0]))
    doc.add_page_break()


def add_results_table(doc, results):
    doc.add_paragraph("สรุปผล Test Case ทั้งหมด", style="Heading 1")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_width(table, [1.15, 3.0, 0.85, 1.0])
    headers = ["Case ID", "Flow / กรณีทดสอบ", "ผล", "หลักฐาน"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
    for item in results:
        row = table.add_row().cells
        row[0].text = item["case_id"]
        row[1].text = item["name"]
        row[2].text = "Passed" if item["passed"] else "Failed"
        row[3].text = item["detail"]
        if not item["passed"]:
            for cell in row:
                set_cell_shading(cell, "FDECEC")
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    set_run_font(run, size=8.8, bold=(row is table.rows[0]))


def friendly_image(path: Path, title: str, subtitle: str) -> Path:
    DOCX_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.open(path).convert("RGB")
    canvas = Image.new("RGB", (1600, 1050), "white")
    image.thumbnail((1504, 900))
    x = (1600 - image.width) // 2
    y = (1050 - image.height) // 2
    canvas.paste(image, (x, y))
    out = DOCX_IMAGE_DIR / path.name
    canvas.save(out, quality=92)
    return out


def add_screenshots(doc, stamp):
    doc.add_page_break()
    doc.add_paragraph("Screen Capture Evidence", style="Heading 1")
    ordered_names = list(SCREENSHOT_LABELS.keys())
    for index, name in enumerate(ordered_names, start=1):
        source = SCREENSHOTS / f"{stamp}-{name}.png"
        if not source.exists():
            continue
        case_id, caption = SCREENSHOT_LABELS[name]
        p = doc.add_paragraph(f"ภาพที่ {index}: {caption}", style="Heading 2")
        p.paragraph_format.keep_with_next = True
        subtitle = f"{case_id} | {source.name}"
        prepared = friendly_image(source, caption, subtitle)
        doc.add_picture(str(prepared), width=Inches(6.4))


def add_notes(doc, result):
    doc.add_page_break()
    doc.add_paragraph("หมายเหตุการทดสอบ", style="Heading 1")
    notes = [
        "Automated backend test และ frontend build ผ่านก่อนทำ live evidence capture",
        "Live API test ครอบคลุม validation, duplicate, staff update, shipping update, tracking lookup, export/import และ permission",
        "UI screenshots จับจาก production URL เพื่อใช้เป็นหลักฐานการทำงานจริง",
        "ข้อมูลทดสอบที่สร้างระหว่าง live test ถูกแยกออกจากงานจริงโดยตั้งสถานะเป็น cancelled",
        f'รายการทดสอบที่สร้างจาก API run: {", ".join(result.get("created_test_requests", []))}',
    ]
    for note in notes:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(note)
        set_run_font(r)


def main():
    result_file = latest_json()
    result = json.loads(result_file.read_text(encoding="utf-8"))
    stamp = latest_screenshot_stamp()

    doc = Document()
    setup_styles(doc)
    add_cover(doc, result)
    add_results_table(doc, result["results"])
    add_screenshots(doc, stamp)
    add_notes(doc, result)
    doc.core_properties.title = "Mahosample Test Evidence With Screenshots"
    doc.core_properties.subject = "รายงานผลทดสอบระบบพร้อมภาพหน้าจอ"
    doc.core_properties.author = "Mahosample Project Team"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
