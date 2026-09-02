from __future__ import annotations

from pathlib import Path

from build_final_docx import add_markdown_file, set_run_font, set_table_width, setup_styles
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUT = DOCS / "Mahosample-staff-user-guide.docx"


def add_staff_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Mahosample")
    set_run_font(r, size=28, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Staff User Guide")
    set_run_font(r, size=18, color="2E74B5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("คู่มือการใช้งานระบบสำหรับพนักงาน")
    set_run_font(r, size=14, color="333333")

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    set_table_width(table, [1.8, 4.6])
    rows = [
        ("Project", "Mahosample"),
        ("Document", "คู่มือการใช้งานสำหรับพนักงาน"),
        ("Live URL", "http://maho.kitaith.com:18080/staff"),
        ("Date", "2026-09-02"),
        ("Purpose", "ใช้อธิบายขั้นตอนการทำงานหลักของพนักงานในระบบ"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(cell is row.cells[0]))

    doc.add_page_break()


def main():
    doc = Document()
    setup_styles(doc)
    add_staff_cover(doc)
    add_markdown_file(doc, "คู่มือการใช้งานสำหรับพนักงาน", DOCS / "staff-user-guide.md", first=True)
    doc.core_properties.title = "Mahosample Staff User Guide"
    doc.core_properties.subject = "คู่มือการใช้งานระบบ Mahosample สำหรับพนักงาน"
    doc.core_properties.author = "Mahosample Project Team"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
