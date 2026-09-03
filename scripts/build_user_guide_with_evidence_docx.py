from __future__ import annotations

import json
import re
from pathlib import Path

from build_final_docx import (
    SCREENSHOT_LABELS,
    SCREENSHOTS,
    add_code_block,
    add_markdown_file,
    add_markdown_text,
    add_page_number,
    add_screenshot_appendix,
    add_table,
    latest_screenshot_stamp,
    prepare_screenshot_for_docx,
    set_run_font,
    set_table_width,
    split_markdown_table,
    setup_styles,
)
from build_test_evidence_docx import add_results_table, latest_json
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUT = DOCS / "Mahosample-user-guide-with-test-cases-and-screenshots.docx"

INLINE_SCREENSHOTS = {
    "1. เข้าใช้งานหน้าแบบฟอร์ม": ["public-registration-form"],
    "2. กรอกข้อมูลผู้ลงทะเบียน": ["public-registration-form", "public-contact-validation"],
    "3. กรอกข้อมูลจัดส่ง": ["public-registration-form"],
    "4. ยินยอมใช้ข้อมูลส่วนบุคคล": ["public-registration-form"],
    "5. ส่งแบบฟอร์มสำเร็จ": ["public-success-card"],
    "6. กรณีข้อมูลซ้ำ": ["public-contact-validation"],
    "7. ตรวจสอบสถานะพัสดุ": ["tracking-found"],
    "8. ไม่พบเลข tracking": ["tracking-not-found"],
    "9. แนวทางความปลอดภัยของข้อมูล": ["public-registration-form"],
    "1. เข้าใช้งาน": ["staff-dashboard"],
    "2. หน้า Dashboard": ["staff-dashboard"],
    "5. ปุ่มบันทึกคำขอ": ["staff-detail-drawer"],
    "6. ปุ่มอัปเดตขนส่ง": ["staff-detail-drawer", "staff-tracking-url-link"],
    "7. หน้าตารางข้อมูลเต็ม": ["staff-full-table"],
    "8. การค้นหาและจัดเรียงในตาราง": [
        "staff-filter-province",
        "staff-date-range-filter",
        "staff-sorting",
    ],
    "9. Tracking URL": ["staff-tracking-url-link"],
    "10. Export CSV สำหรับไปรษณีย์": ["staff-export"],
    "11. Import tracking กลับเข้าระบบ": ["staff-full-table"],
    "12. เช็คสถานะให้ลูกค้า": ["tracking-found"],
    "14. ปัญหาที่พบบ่อย": ["staff-full-table"],
    "TC-PUB-001 ส่งแบบฟอร์มลงทะเบียนสำเร็จ": ["public-registration-form", "public-success-card"],
    "TC-PUB-002 ตรวจ validation เมื่อไม่กรอกช่องทางติดต่อ": ["public-contact-validation"],
    "TC-PUB-003 ตรวจ validation เบอร์โทรศัพท์": ["public-contact-validation"],
    "TC-PUB-004 ตรวจ validation รหัสไปรษณีย์": ["public-contact-validation"],
    "TC-PUB-005 ตรวจ validation PDPA": ["public-contact-validation"],
    "TC-PUB-006 ตรวจ duplicate เบอร์โทรศัพท์": ["public-contact-validation"],
    "TC-PUB-007 ตรวจ duplicate ที่อยู่จัดส่ง": ["public-contact-validation"],
    "TC-PUB-008 ตรวจ duplicate ทั้งเบอร์โทรศัพท์และที่อยู่จัดส่ง": ["public-contact-validation"],
    "TC-TRK-001 เช็คสถานะด้วยเลข tracking": ["tracking-found"],
    "TC-TRK-002 ไม่พบเลข tracking": ["tracking-not-found"],
    "TC-STF-001 Login พนักงานสำเร็จ": ["staff-dashboard"],
    "TC-STF-003 ดู Dashboard พนักงาน": ["staff-dashboard"],
    "TC-STF-004 บันทึกคำขอ": ["staff-detail-drawer"],
    "TC-STF-005 อัปเดตขนส่ง": ["staff-detail-drawer"],
    "TC-STF-006 ตารางข้อมูลเต็ม แสดงข้อมูลและเปิดรายละเอียด": [
        "staff-full-table",
        "staff-detail-drawer",
    ],
    "TC-STF-007 Sorting ใน header ตาราง": ["staff-sorting"],
    "TC-STF-008 Filter ใน header ตาราง": ["staff-filter-province"],
    "TC-STF-009 Filter วันที่ส่งแบบช่วงเวลา": ["staff-date-range-filter"],
    "TC-STF-010 Filter วันที่ลงทะเบียนแบบช่วงเวลา": ["staff-date-range-filter"],
    "TC-STF-011 แสดงผลครั้งละ 100 รายการ": ["staff-full-table"],
    "TC-STF-012 Tracking URL คลิกได้": ["staff-tracking-url-link"],
    "TC-EXP-001 Export CSV สำหรับไปรษณีย์": ["staff-export"],
    "TC-EXP-002 Export CSV จากตารางข้อมูลเต็มตาม filter": ["staff-export"],
    "TC-IMP-001 Import tracking จาก CSV": ["staff-full-table"],
    "TC-IMP-002 Import tracking จากไฟล์ KEX XLSX": ["staff-full-table"],
}


def setup_user_guide_styles(doc):
    setup_styles(doc)
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.clear()
    set_run_font(footer.add_run("Mahosample User Guide | Page "), size=9, color="666666")
    add_page_number(footer)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Mahosample")
    set_run_font(r, size=28, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("User Guide And Test Evidence")
    set_run_font(r, size=18, color="2E74B5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("คู่มือการใช้งานฝั่งลูกค้าและพนักงาน พร้อม Test Cases และภาพหน้าจอประกอบ")
    set_run_font(r, size=14, color="333333")

    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=2)
    set_table_width(table, [1.9, 4.5])
    rows = [
        ("Project", "Mahosample"),
        ("Document", "User Guide With Test Cases And Screen Capture Evidence"),
        ("Public URL", "http://maho.kitaith.com:18080/"),
        ("Staff URL", "http://maho.kitaith.com:18080/staff"),
        ("Tracking URL", "http://maho.kitaith.com:18080/tracking"),
        ("Date", "2026-09-03"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(cell is row.cells[0]))

    doc.add_page_break()


def add_overview(doc):
    doc.add_paragraph("ภาพรวมเอกสาร", style="Heading 1")
    intro = [
        "เอกสารฉบับนี้แยกมาจากส่วนคู่มือและหลักฐานทดสอบของ Mahosample-final-project-report.docx เพื่อใช้เป็นคู่มือปฏิบัติงานและหลักฐานการใช้งานระบบโดยตรง",
        "เนื้อหาครอบคลุมฝั่งลูกค้าทั่วไป ฝั่งพนักงาน รายการ Test Cases ทั้งหมด และภาพหน้าจอจาก production environment",
    ]
    for text in intro:
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run_font(r)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [1.6, 3.0, 1.8])
    headers = ["ส่วน", "เนื้อหา", "ผู้ใช้หลัก"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    rows = [
        ("Customer Guide", "ขั้นตอนลงทะเบียนและเช็คสถานะ tracking", "ลูกค้าทั่วไป"),
        ("Staff Guide", "Dashboard, full table, บันทึกคำขอ, อัปเดตขนส่ง, import/export", "พนักงาน"),
        ("Test Cases", "กรณีทดสอบทุก flow หลัก", "ทีมตรวจรับ/ทีมพัฒนา"),
        ("Screen Capture Evidence", "ภาพหน้าจอประกอบการทำงานจริงบน production", "ทีมตรวจรับ"),
    ]
    for values in rows:
        row = table.add_row().cells
        for cell, value in zip(row, values):
            cell.text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, bold=(row is table.rows[0]))


def add_inline_screenshots(doc, heading_text):
    stamp = latest_screenshot_stamp()
    if not stamp:
        return

    for name in INLINE_SCREENSHOTS.get(heading_text, []):
        source = SCREENSHOTS / f"{stamp}-{name}.png"
        if not source.exists():
            continue
        caption = SCREENSHOT_LABELS.get(name, name)
        p = doc.add_paragraph(f"ภาพประกอบ: {caption}", style="Heading 3")
        p.paragraph_format.keep_with_next = True
        prepared = prepare_screenshot_for_docx(source)
        doc.add_picture(str(prepared), width=Inches(5.9))


def add_markdown_file_with_inline_screenshots(doc, title, path, first=False):
    if not first:
        doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph(title, style="Heading 1")
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            rows, i = split_markdown_table(lines, i)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            heading_text = heading.group(2).strip()
            p = doc.add_paragraph(heading_text, style=f"Heading {level}")
            p.paragraph_format.keep_with_next = bool(INLINE_SCREENSHOTS.get(heading_text))
            add_inline_screenshots(doc, heading_text)
            i += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_markdown_text(p, bullet.group(1))
            i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_markdown_text(p, numbered.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_markdown_text(p, line)
        i += 1

    if in_code and code_lines:
        add_code_block(doc, "\n".join(code_lines))


def add_test_summary(doc):
    result_file = latest_json()
    result = json.loads(result_file.read_text(encoding="utf-8"))
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("สรุปผลการทดสอบ", style="Heading 1")
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            f'Live E2E run {result["run_id"]}: ผ่าน {result["passed"]}/{result["total"]} เคส, failed {result["failed"]} เคส',
        ),
    )
    add_results_table(doc, result["results"])


def main():
    doc = Document()
    setup_user_guide_styles(doc)
    add_cover(doc)
    add_overview(doc)
    add_markdown_file_with_inline_screenshots(doc, "คู่มือฝั่งลูกค้า", DOCS / "public-user-guide.md")
    add_markdown_file_with_inline_screenshots(doc, "คู่มือฝั่งพนักงาน", DOCS / "staff-user-guide.md")
    add_markdown_file_with_inline_screenshots(doc, "Test Cases ทั้งหมด", DOCS / "test-cases.md")
    add_test_summary(doc)
    add_screenshot_appendix(doc)
    doc.core_properties.title = "Mahosample User Guide With Test Cases And Screenshots"
    doc.core_properties.subject = "คู่มือการใช้งานฝั่งลูกค้าและพนักงาน พร้อม Test Cases และภาพหน้าจอ"
    doc.core_properties.author = "Mahosample Project Team"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
